#!/usr/bin/env python3
"""
DOC (.doc) → DOCX → TXT (Markdown, LEGAL AWARE)

- Convert .doc → .docx via LibreOffice (headless)
- Extract tables with python-docx
- Detect merged-cell section headers
- Chunk output = Title + Annex + Header + Data
"""

import argparse
import subprocess
from pathlib import Path
from docx import Document

# ==========================================================
# CONFIG
# ==========================================================

DOCUMENT_TITLE = (
    """Quyết định 1227/QĐ-BYT năm 2025 về Danh mục mã dùng chung đối với kỹ thuật, thuật ngữ chỉ số cận lâm sàng (Đợt 1) do Bộ trưởng Bộ Y tế ban hành"""
)

CHUNK_ROW_SIZE = 30

# ==========================================================
# UTILITIES
# ==========================================================

def normalize_text(text: str) -> str:
    if not text:
        return ""
    return " ".join(
        text.replace("\r", " ")
            .replace("\n", " ")
            .replace("\t", " ")
            .split()
    )


def chunk_rows(rows, size):
    for i in range(0, len(rows), size):
        yield rows[i:i + size]

# ==========================================================
# DOC → DOCX
# ==========================================================

def convert_doc_to_docx(doc_path: Path, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)

    before = set(out_dir.glob("*.docx"))

    subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "docx",
            str(doc_path),
            "--outdir",
            str(out_dir),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )

    after = set(out_dir.glob("*.docx"))
    created = after - before

    if not created:
        raise RuntimeError(f"❌ LibreOffice convert failed: {doc_path.name}")

    # luôn chỉ có 1 file được tạo
    return created.pop()

# ==========================================================
# SECTION HEADER DETECTION (KEY FIX)
# ==========================================================

def is_section_header_row(row: list[str]) -> bool:
    """
    Detect merged-cell rows where all non-empty cells are identical.
    """
    texts = [c for c in row if c.strip()]
    if not texts:
        return False

    if len(set(texts)) != 1:
        return False

    t = texts[0].upper()
    return (
        "XÉT NGHIỆM" in t
        or t.startswith(("I.", "II.", "III.", "1.", "2.", "3.", "CHƯƠNG", "MỤC"))
    )


def extract_section_title(row: list[str]) -> str | None:
    for c in row:
        if c.strip():
            return c.strip()
    return None

# ==========================================================
# HEADER DETECTION
# ==========================================================

def detect_header_row(rows, scan=10):
    for i, row in enumerate(rows[:scan]):
        for c in row:
            t = c.upper()
            if "STT" in t or "MÃ" in t:
                return i
    return None

# ==========================================================
# ANNEX TITLE
# ==========================================================

def extract_annex_title(paragraphs, limit=12):
    lines = []
    for p in paragraphs[:limit]:
        t = normalize_text(p.text)
        if t:
            lines.append(t)
    text = " ".join(lines)
    return text if len(text) > 40 else None

# ==========================================================
# WRITE CHUNK
# ==========================================================

def write_chunk(f, headers, rows, annex, section):
    f.write(DOCUMENT_TITLE + "\n")
    if annex:
        f.write(annex + "\n")
    if section:
        f.write(f"### {section}\n")

    f.write("| " + " | ".join(headers) + " |\n")
    f.write("|" + "|".join(["---"] * len(headers)) + "|\n")
    for r in rows:
        f.write("| " + " | ".join(r) + " |\n")
    f.write("\n")

# ==========================================================
# MAIN PIPELINE
# ==========================================================

def process_folder(input_dir, output_dir):
    input_dir = Path(input_dir)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    tmp_docx = output_dir / "_docx"
    tmp_docx.mkdir(exist_ok=True)

    for doc_file in input_dir.glob("*.doc"):
        print(f"📄 {doc_file.name}")

        docx_file = convert_doc_to_docx(doc_file, tmp_docx)
        doc = Document(docx_file)

        annex = extract_annex_title(doc.paragraphs)
        out_txt = output_dir / f"{doc_file.stem}.txt"

        with open(out_txt, "w", encoding="utf-8") as f:
            for table in doc.tables:
                rows = [
                    [normalize_text(c.text) for c in r.cells]
                    for r in table.rows
                ]

                header_idx = detect_header_row(rows)
                if header_idx is None:
                    continue

                headers = rows[header_idx]
                data = rows[header_idx + 1:]

                section = None
                buffer = []

                for row in data:
                    if is_section_header_row(row):
                        if buffer:
                            for chunk in chunk_rows(buffer, CHUNK_ROW_SIZE):
                                write_chunk(f, headers, chunk, annex, section)
                            buffer = []
                        section = extract_section_title(row)
                    elif any(c.strip() for c in row):
                        buffer.append(row)

                if buffer:
                    for chunk in chunk_rows(buffer, CHUNK_ROW_SIZE):
                        write_chunk(f, headers, chunk, annex, section)

        print(f"✅ Saved: {out_txt}")

# ==========================================================
# CLI
# ==========================================================

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input", help="Folder chứa file .doc")
    ap.add_argument("-o", "--output", default="output_txt")
    args = ap.parse_args()

    process_folder(args.input, args.output)

if __name__ == "__main__":
    main()
